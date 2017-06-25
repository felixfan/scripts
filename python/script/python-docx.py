#coding=utf-8
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# Opening a document
document = Document()

# Adding a heading
'''
In anything but the shortest document,
body text is divided into sections,
each of which starts with a heading.
Here’s how to add one, by default: level=1
'''
document.add_heading('The REAL meaning of the universe')
document.add_heading('The role of dolphins', level=2)

# Adding a paragraph
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

'''
When you add a paragraph by providing text to the .add_paragraph() method, it gets put into a single run. You can add more using the
.add_run() method on the paragraph:
'''
paragraph = document.add_paragraph('Add more text ')
paragraph.add_run('by .add_run() method.')

# Applying bold and italic
paragraph = document.add_paragraph('Applying ')
paragraph.add_run('bold ').bold = True
paragraph.add_run('and ')
paragraph.add_run('italic.').italic = True

# Applying a character style
paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.').style = 'Emphasis'


# Applying a paragraph style
document.add_paragraph('This particular style causes the paragraph to appear as a bullet', style='List Bullet')
document.add_paragraph('This particular style causes the paragraph to appear as a bullet', style='List Number')

# Indentation
s = 'Indentation is the horizontal space between a paragraph and edge of its container, typically the page margin. The first line can also have a different indentation than the rest of the paragraph.'
paragraph = document.add_paragraph(s)
paragraph.paragraph_format.first_line_indent = Inches(0.25)
# Line spacing
paragraph.paragraph_format.line_spacing = 1.75

# Apply character formatting
run1 = document.add_paragraph().add_run('Character formatting is applied at the Run level. Examples include font typeface and size, bold, italic, and underline.')
font1 = run1.font
font1.name = 'Calibri'
font1.size = Pt(12)
font1.italic = True
font1.underline = True
font1.bold = True
font1.color.rgb = RGBColor(250, 0, 0)
'''
black (0,0,0)
white (255, 255, 255)
red   (255, 0, 0)
green (0, 255, 0)
blue  (0, 0, 255)
'''

# Adding a table
nrow = 4
ncol = 3
table = document.add_table(rows=nrow, cols=ncol)
table.style = "Light Grid Accent 2"
info = [
    ['ID', 'Height', 'Weight'],
    ['a', 170, 60],
    ['b', 160, 45],
    ['c', 176, 66]
    ]
for i in range(nrow):
    for j in range(ncol):
        table.cell(i, j).text = str(info[i][j])

# Adding a page break
# document.add_page_break()

# Adding a picture
document.add_picture('hku.jpg', width=Inches(5.0))

document.save('demo.docx')
